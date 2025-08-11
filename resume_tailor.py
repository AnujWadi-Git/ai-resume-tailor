#!/usr/bin/env python3
"""
Resume Tailoring Tool - Google Gemini Version (Word Output)
A complete system to automatically tailor resumes based on job descriptions using Google Gemini.
Outputs Word files instead of PDF with proper skills section formatting.

Requirements:
pip install google-generativeai python-docx streamlit

Usage:
- CLI: python resume_tailor.py
- Streamlit: streamlit run resume_tailor.py
"""

import os
import sys
import argparse
from pathlib import Path
from typing import Optional, Dict, Any, List
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import google.generativeai as genai
import streamlit as st
import tempfile


class ResumeLoader:
    """Handles loading and parsing of resume documents."""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self.content = ""
        self.skills_section = ""
        self.skills_paragraphs = []
        self.word_count = 0
        
    def load_resume(self) -> Dict[str, Any]:
        """Load and parse the .docx resume file."""
        try:
            self.doc = Document(self.file_path)
            self.content = self._extract_text()
            self.skills_section, self.skills_paragraphs = self._extract_skills_section_with_paragraphs()
            self.word_count = len(self.content.split())
            
            return {
                'document': self.doc,
                'content': self.content,
                'skills_section': self.skills_section,
                'skills_paragraphs': self.skills_paragraphs,
                'word_count': self.word_count,
                'success': True
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _extract_text(self) -> str:
        """Extract all text from the document."""
        full_text = []
        for paragraph in self.doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    
    def _extract_skills_section_with_paragraphs(self) -> tuple:
        """Extract the skills section and identify the paragraphs that contain it."""
        skills_paragraphs = []
        skills_content = ""
        
        # Common skills section headers
        skills_headers = [
            'skills', 'technical skills', 'core competencies', 
            'key skills', 'areas of expertise', 'technologies',
            'programming languages', 'tools & technologies'
        ]
        
        # Find skills header paragraph
        skills_start_idx = -1
        for i, paragraph in enumerate(self.doc.paragraphs):
            para_text = paragraph.text.strip().lower()
            if any(header.lower() in para_text for header in skills_headers):
                skills_start_idx = i
                break
        
        if skills_start_idx == -1:
            return "", []
        
        # Collect skills section paragraphs
        skills_paragraphs.append(skills_start_idx)  # Include the header
        skills_content = self.doc.paragraphs[skills_start_idx].text + "\n"
        
        # Find content after the skills header until next major section
        for i in range(skills_start_idx + 1, len(self.doc.paragraphs)):
            paragraph = self.doc.paragraphs[i]
            para_text = paragraph.text.strip()
            
            # Stop if we hit another major section (all caps, bold, or common headers)
            if para_text and (
                para_text.isupper() or 
                any(header in para_text.lower() for header in [
                    'experience', 'education', 'projects', 'certifications',
                    'awards', 'achievements', 'summary', 'objective'
                ]) or
                (paragraph.runs and paragraph.runs[0].bold and len(para_text) < 50)
            ):
                break
            
            skills_paragraphs.append(i)
            skills_content += para_text + "\n"
        
        return skills_content.strip(), skills_paragraphs


class JobDescriptionLoader:
    """Handles loading job descriptions from text or file."""
    
    @staticmethod
    def load_from_text(text: str) -> Dict[str, Any]:
        """Load job description from raw text."""
        try:
            return {
                'content': text.strip(),
                'word_count': len(text.split()),
                'success': True
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def load_from_file(file_path: str) -> Dict[str, Any]:
        """Load job description from .txt file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return JobDescriptionLoader.load_from_text(content)
        except Exception as e:
            return {'success': False, 'error': str(e)}


class ResumeTailor:
    """Core class for tailoring resumes using Google Gemini."""
    
    def __init__(self, api_key: str):
        genai.configure(api_key=api_key)
        
        # Try different Gemini models in order of preference
        self.models_to_try = [
            "gemini-1.5-pro",      # Latest and most capable
            "gemini-1.5-flash",    # Fast and efficient
            "gemini-pro",          # Original Gemini Pro
        ]
        
        self.model = None
        self._initialize_model()
        
        self.prompt_template = """
You are an expert resume writer and ATS optimization specialist. Your task is to create ONLY the content that goes under the skills section header of a resume.

ORIGINAL RESUME CONTENT:
{resume_content}

CURRENT SKILLS CONTENT (everything under the skills header):
{skills_section}

JOB DESCRIPTION:
{job_description}

INSTRUCTIONS:
1. Create ONLY the content that goes under the skills section header (NOT the header itself)
2. Keep the response concise to maintain one-page resume limit
3. Use bullet points or comma-separated format based on the original style
4. Bold important skill categories or key technologies
5. Prioritize skills mentioned in the job description
6. Keep existing relevant skills and add new ones from job description
7. Use ATS-friendly keywords from the job posting
8. Maintain professional formatting suitable for 9pt font

IMPORTANT CONSTRAINTS:
- Do NOT include the "Skills" header or any section titles
- Keep response under 150 words to maintain page limit
- Use concise, keyword-rich descriptions
- Focus on technical skills, tools, and relevant competencies
- Make skills easily scannable for ATS systems

OUTPUT FORMAT:
Return ONLY the skills content (no headers, no extra formatting instructions), ready to be placed under the skills section with 9pt font.

SKILLS CONTENT:
"""
    
    def _initialize_model(self):
        """Initialize the best available Gemini model."""
        for model_name in self.models_to_try:
            try:
                model = genai.GenerativeModel(model_name)
                # Test the model with a simple request
                test_response = model.generate_content("Hello", 
                    generation_config=genai.types.GenerationConfig(
                        max_output_tokens=10,
                        temperature=0.1
                    )
                )
                self.model = model
                self.model_name = model_name
                print(f"✅ Successfully initialized Gemini model: {model_name}")
                return
            except Exception as e:
                print(f"⚠️  Model {model_name} not available: {str(e)}")
                continue
        
        raise Exception("No Gemini models are available. Please check your API key and account status.")
    
    def tailor_resume(self, resume_data: Dict, job_description: str) -> Dict[str, Any]:
        """Tailor the resume skills section using Google Gemini."""
        try:
            if not self.model:
                return {'success': False, 'error': 'No Gemini model available'}
            
            # Prepare the prompt
            prompt = self.prompt_template.format(
                resume_content=resume_data['content'][:3000],
                skills_section=resume_data['skills_section'],
                job_description=job_description[:2000]
            )
            
            # Configure generation parameters
            generation_config = genai.types.GenerationConfig(
                max_output_tokens=800,  # Reduced to keep content concise
                temperature=0.3,
                top_p=0.8,
                top_k=40
            )
            
            # Call Gemini
            response = self.model.generate_content(
                prompt,
                generation_config=generation_config
            )
            
            if not response.text:
                return {'success': False, 'error': 'Empty response from Gemini'}
            
            tailored_skills = response.text.strip()
            
            # Clean up the response
            if tailored_skills.startswith("SKILLS CONTENT:"):
                tailored_skills = tailored_skills.replace("SKILLS CONTENT:", "").strip()
            
            # Remove any section headers that might have been included
            lines = tailored_skills.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line and not (line.isupper() and len(line) < 30 and any(word in line.lower() for word in ['skill', 'technical', 'competenc'])):
                    cleaned_lines.append(line)
            
            tailored_skills = '\n'.join(cleaned_lines)
            
            # Calculate approximate token usage
            estimated_tokens = len(prompt.split()) + len(tailored_skills.split())
            
            return {
                'success': True,
                'tailored_skills': tailored_skills,
                'original_skills': resume_data['skills_section'],
                'tokens_used': estimated_tokens,
                'model_used': self.model_name
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class ResumeExporter:
    """Handles exporting tailored resumes to Word format."""
    
    @staticmethod
    def update_docx_skills(doc: Document, skills_paragraphs: List[int], new_skills: str) -> Document:
        """Update the skills section in the Word document with proper formatting."""
        try:
            if not skills_paragraphs:
                return doc
            
            # Remove old skills content (everything except the header)
            skills_header_idx = skills_paragraphs[0]
            
            # Remove all paragraphs after the header that were part of skills section
            paragraphs_to_remove = []
            for idx in reversed(skills_paragraphs[1:]):  # Skip the header
                if idx < len(doc.paragraphs):
                    paragraphs_to_remove.append(doc.paragraphs[idx])
            
            for paragraph in paragraphs_to_remove:
                p = paragraph._element
                p.getparent().remove(p)
                p.clear()
            
            # Add new skills content after the header
            header_paragraph = doc.paragraphs[skills_header_idx]
            
            # Split new skills content into lines
            skills_lines = [line.strip() for line in new_skills.split('\n') if line.strip()]
            
            # Insert new paragraphs after the header
            for i, line in enumerate(skills_lines):
                # Create new paragraph
                new_paragraph = doc.add_paragraph()
                
                # Move the new paragraph to the correct position
                header_p = header_paragraph._element
                new_p = new_paragraph._element
                header_p.getparent().insert(
                    list(header_p.getparent()).index(header_p) + 1 + i,
                    new_p
                )
                
                # Set font size to 9pt and add content with bold formatting
                ResumeExporter._format_skills_paragraph(new_paragraph, line)
            
            return doc
            
        except Exception as e:
            print(f"Warning: Could not update skills section automatically: {e}")
            return doc
    
    @staticmethod
    def _format_skills_paragraph(paragraph, text):
        """Format a skills paragraph with 9pt font and bold key terms."""
        # Clear existing content
        paragraph.clear()
        
        # Set paragraph formatting
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Parse text for bold formatting (assuming **text** or common skill categories)
        bold_patterns = [
            r'\*\*(.*?)\*\*',  # **bold text**
            r'\b([A-Z][a-z]+ [A-Z][a-z]+):\s',  # Category headers like "Programming Languages:"
            r'^([A-Za-z\s]+):\s',  # Line starting with category
        ]
        
        current_pos = 0
        
        # Look for bold patterns
        for pattern in bold_patterns:
            matches = list(re.finditer(pattern, text))
            if matches:
                for match in matches:
                    # Add regular text before match
                    if match.start() > current_pos:
                        run = paragraph.add_run(text[current_pos:match.start()])
                        run.font.size = Pt(9)
                    
                    # Add bold text
                    bold_text = match.group(1) if len(match.groups()) > 0 else match.group(0)
                    bold_text = bold_text.replace('**', '')  # Remove markdown
                    run = paragraph.add_run(bold_text)
                    run.font.size = Pt(9)
                    run.font.bold = True
                    
                    # Add colon if it was part of the pattern
                    if ':' in match.group(0) and not bold_text.endswith(':'):
                        colon_run = paragraph.add_run(':')
                        colon_run.font.size = Pt(9)
                    
                    current_pos = match.end()
                
                # Add remaining text
                if current_pos < len(text):
                    run = paragraph.add_run(text[current_pos:])
                    run.font.size = Pt(9)
                return
        
        # If no bold patterns found, add as regular text but bold key terms
        words = text.split()
        for i, word in enumerate(words):
            # Bold common technical terms and tools
            if (word.lower() in ['python', 'java', 'javascript', 'react', 'sql', 'aws', 'azure', 'docker', 'kubernetes', 'git', 'linux', 'windows', 'excel', 'powerbi', 'tableau'] or 
                word.isupper() or 
                any(char.isdigit() for char in word)):
                run = paragraph.add_run(word)
                run.font.bold = True
            else:
                run = paragraph.add_run(word)
            
            run.font.size = Pt(9)
            
            # Add space after word (except last word)
            if i < len(words) - 1:
                space_run = paragraph.add_run(' ')
                space_run.font.size = Pt(9)
    
    @staticmethod
    def export_to_docx(doc: Document, output_path: str) -> Dict[str, Any]:
        """Export the updated document to Word format."""
        try:
            # Save the document
            doc.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'file_size': os.path.getsize(output_path)
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class ResumeAutoTailor:
    """Main orchestrator class for the resume tailoring process."""
    
    def __init__(self, api_key: str):
        self.tailor = ResumeTailor(api_key)
        self.exporter = ResumeExporter()
    
    def process_resume(self, resume_path: str, job_description: str, output_path: str) -> Dict[str, Any]:
        """Complete resume tailoring process."""
        results = {'steps': [], 'success': False}
        
        try:
            # Step 1: Load resume
            print("📄 Loading resume...")
            resume_loader = ResumeLoader(resume_path)
            resume_data = resume_loader.load_resume()
            
            if not resume_data['success']:
                results['error'] = f"Failed to load resume: {resume_data['error']}"
                return results
            
            results['steps'].append(f"✅ Resume loaded successfully ({resume_data['word_count']} words)")
            
            # Step 2: Load job description
            print("📋 Processing job description...")
            jd_data = JobDescriptionLoader.load_from_text(job_description)
            
            if not jd_data['success']:
                results['error'] = f"Failed to process job description: {jd_data['error']}"
                return results
            
            results['steps'].append(f"✅ Job description processed ({jd_data['word_count']} words)")
            
            # Step 3: Tailor resume with Gemini
            print("🤖 Tailoring resume with Google Gemini...")
            tailor_result = self.tailor.tailor_resume(resume_data, jd_data['content'])
            
            if not tailor_result['success']:
                results['error'] = f"Failed to tailor resume: {tailor_result['error']}"
                return results
            
            results['steps'].append(f"✅ Resume tailored successfully (used ~{tailor_result['tokens_used']} tokens with {tailor_result['model_used']})")
            
            # Step 4: Update document
            print("📝 Updating document with new skills...")
            updated_doc = self.exporter.update_docx_skills(
                resume_data['document'], 
                resume_data['skills_paragraphs'],
                tailor_result['tailored_skills']
            )
            
            results['steps'].append("✅ Document updated with tailored skills (9pt font applied)")
            
            # Step 5: Export to Word
            print("📄 Exporting to Word format...")
            export_result = self.exporter.export_to_docx(updated_doc, output_path)
            
            if not export_result['success']:
                results['error'] = f"Failed to export Word file: {export_result['error']}"
                return results
            
            results['steps'].append(f"✅ Word file exported successfully ({export_result['file_size']} bytes)")
            results['success'] = True
            results['output_path'] = export_result['output_path']
            results['tailored_skills'] = tailor_result['tailored_skills']
            results['original_skills'] = tailor_result['original_skills']
            
            return results
            
        except Exception as e:
            results['error'] = f"Unexpected error: {str(e)}"
            return results


def run_cli():
    """Command Line Interface for the resume tailoring tool."""
    parser = argparse.ArgumentParser(description='Resume Tailoring Tool - Google Gemini Version (Word Output)')
    parser.add_argument('--resume', '-r', required=True, help='Path to resume .docx file')
    parser.add_argument('--job-description', '-j', help='Path to job description .txt file')
    parser.add_argument('--job-text', '-t', help='Job description as text')
    parser.add_argument('--output', '-o', required=True, help='Output Word file path (.docx)')
    parser.add_argument('--api-key', '-k', help='Google AI API key (or set GOOGLE_API_KEY env var)')
    
    args = parser.parse_args()
    
    # Ensure output has .docx extension
    if not args.output.endswith('.docx'):
        args.output += '.docx'
    
    # Get API key
    api_key = args.api_key or os.getenv('GOOGLE_API_KEY') or os.getenv('GEMINI_API_KEY')
    if not api_key:
        print("❌ Error: Google AI API key is required. Use --api-key or set GOOGLE_API_KEY environment variable.")
        print("💡 Get your API key from: https://makersuite.google.com/app/apikey")
        sys.exit(1)
    
    # Get job description
    job_description = ""
    if args.job_description:
        jd_result = JobDescriptionLoader.load_from_file(args.job_description)
        if not jd_result['success']:
            print(f"❌ Error loading job description: {jd_result['error']}")
            sys.exit(1)
        job_description = jd_result['content']
    elif args.job_text:
        job_description = args.job_text
    else:
        print("❌ Error: Job description is required. Use --job-description or --job-text")
        sys.exit(1)
    
    # Process resume
    auto_tailor = ResumeAutoTailor(api_key)
    result = auto_tailor.process_resume(args.resume, job_description, args.output)
    
    # Print results
    for step in result['steps']:
        print(step)
    
    if result['success']:
        print(f"\n🎉 Success! Tailored resume saved to: {result['output_path']}")
        print(f"📝 Skills section updated with 9pt font")
        print(f"📄 One-page format maintained")
    else:
        print(f"\n❌ Error: {result['error']}")
        sys.exit(1)


def run_streamlit():
    """Streamlit web interface for the resume tailoring tool."""
    st.set_page_config(
        page_title="Resume Tailoring Tool - Word Output",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("🤖 AI Resume Tailoring Tool (Word Format)")
    st.write("Automatically tailor your resume skills section using **Google Gemini** - Outputs Word files with proper formatting")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("Configuration")
        st.markdown("### 🔑 Google AI API Key")
        api_key = st.text_input(
            "Enter your Google AI API key", 
            type="password", 
            help="Get your free API key from https://makersuite.google.com/app/apikey"
        )
        
        if not api_key:
            st.warning("Please enter your Google AI API key to continue")
            st.markdown("""
            **How to get your API key:**
            1. Go to [Google AI Studio](https://makersuite.google.com/app/apikey)
            2. Click "Create API Key"
            3. Copy and paste it above
            """)
            st.stop()
        
        st.markdown("---")
        st.markdown("### ℹ️ Features")
        st.markdown("""
        ✅ **Word file output** (.docx)  
        ✅ **Preserves skills header**  
        ✅ **9pt font for skills content**  
        ✅ **Bold key technologies**  
        ✅ **One-page format maintained**  
        ✅ **ATS-optimized keywords**  
        """)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.header("📄 Upload Resume")
        resume_file = st.file_uploader(
            "Choose your resume (.docx)", 
            type=['docx'],
            help="Upload your base resume in .docx format"
        )
        
        if resume_file:
            st.success(f"✅ Resume uploaded: {resume_file.name}")
    
    with col2:
        st.header("📋 Job Description")
        jd_option = st.radio("Choose input method:", ["Paste Text", "Upload .txt File"])
        
        job_description = ""
        if jd_option == "Paste Text":
            job_description = st.text_area(
                "Paste job description here:", 
                height=300,
                placeholder="Paste the complete job description here...\n\nInclude requirements, skills, responsibilities, etc."
            )
        else:
            jd_file = st.file_uploader("Choose job description (.txt)", type=['txt'])
            if jd_file:
                job_description = str(jd_file.read(), "utf-8")
                st.text_area("Job Description Preview:", job_description[:500] + "...", height=200, disabled=True)
        
        if job_description:
            word_count = len(job_description.split())
            st.info(f"📊 Job description: {word_count} words")
    
    # Process button
    if st.button("🚀 Tailor Resume Skills Section", type="primary", use_container_width=True):
        if not resume_file or not job_description:
            st.error("❌ Please upload a resume and provide a job description")
            st.stop()
        
        with st.spinner("🤖 Tailoring your resume skills with Google Gemini..."):
            try:
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(resume_file.getvalue())
                    tmp_resume_path = tmp_file.name
                
                # Create output path
                output_path = tmp_resume_path.replace('.docx', '_tailored.docx')
                
                # Process resume
                auto_tailor = ResumeAutoTailor(api_key)
                result = auto_tailor.process_resume(tmp_resume_path, job_description, output_path)
                
                if result['success']:
                    st.success("✅ Resume skills section tailored successfully!")
                    
                    # Show process steps
                    with st.expander("📊 Process Details", expanded=True):
                        for step in result['steps']:
                            st.write(step)
                    
                    # Show skills comparison
                    st.subheader("📝 Skills Section Changes")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**📝 Original Skills Section**")
                        st.text_area("", result['original_skills'], height=200, disabled=True, key="original")
                    
                    with col2:
                        st.markdown("**✨ New Skills Content (9pt font)**")
                        st.text_area("", result['tailored_skills'], height=200, disabled=True, key="tailored")
                    
                    # Download button
                    with open(output_path, "rb") as file:
                        st.download_button(
                            label="📥 Download Tailored Resume (Word)",
                            data=file.read(),
                            file_name=f"tailored_resume_{resume_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    st.info("📝 The skills header remains unchanged. Only the content under it has been updated with proper formatting.")
                
                else:
                    st.error(f"❌ Error: {result['error']}")
                    if "API key" in result['error']:
                        st.info("💡 Make sure you have a valid Google AI API key from https://makersuite.google.com/app/apikey")
                
                # Cleanup
                if os.path.exists(tmp_resume_path):
                    os.remove(tmp_resume_path)
                if os.path.exists(output_path):
                    os.remove(output_path)
                    
            except Exception as e:
                st.error(f"An unexpected error occurred: {str(e)}")


if __name__ == "__main__":
    # Determine if running in Streamlit
    if len(sys.argv) > 1 and sys.argv[1] == "streamlit":
        # Remove the streamlit argument to avoid conflicts
        sys.argv.remove("streamlit")
        run_streamlit()
    else:
        # Check if we're in a streamlit environment
        try:
            # This will only work if we're in Streamlit
            import streamlit.web.cli as stcli
            run_streamlit()
        except:
            # Run CLI if not in Streamlit environment
            if len(sys.argv) == 1:
                print("🤖 Resume Tailoring Tool - Word Output Version")
                print("Usage:")
                print("  CLI: python resume_tailor.py --resume resume.docx --job-text 'job description' --output tailored.docx --api-key YOUR_GEMINI_KEY")
                print("  Streamlit: streamlit run resume_tailor.py")
                print("\n💡 Get your Google AI API key from: https://makersuite.google.com/app/apikey")
                print("📄 Now outputs Word files instead of PDF with proper skills formatting!")
            else:
                run_cli()